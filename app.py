#!/usr/bin/env python3
"""
HỆ THỐNG PHÂN TÍCH KỸ THUẬT CHỨNG KHOÁN VIỆT NAM
Version: 3.0 - Full Features (Giống Google Colab)
"""

import os
import json
import warnings
import zipfile
from datetime import datetime, timedelta
from pathlib import Path

import numpy as np
import pandas as pd
import gradio as gr
import plotly.graph_objects as go
from plotly.subplots import make_subplots

warnings.filterwarnings('ignore')

# ============================================================
# CẤU HÌNH
# ============================================================

DEFAULT_CONFIG = {
    "watchlist": ["VNM", "FPT", "VIC", "HPG", "MWG", "TCB", "VCB", "ACB", "VPB", "DGW", "VCK"],
    "days_back": 90,
    "min_data_points": 5,
    "output_folder": "./output",
}

CONFIG_FILE = "config.json"

def load_config():
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    save_config(DEFAULT_CONFIG)
    return DEFAULT_CONFIG

def save_config(config):
    with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
        json.dump(config, f, indent=2, ensure_ascii=False)

# ============================================================
# DANH SÁCH 26 CHỈ BÁO
# ============================================================

INDICATOR_INFO = {
    'SMA': {'group': 'trend', 'name': 'SMA (5,10,20,50,100,200)'},
    'EMA': {'group': 'trend', 'name': 'EMA (12,26,50)'},
    'WMA': {'group': 'trend', 'name': 'WMA (10,20)'},
    'TEMA': {'group': 'trend', 'name': 'TEMA (20)'},
    'DEMA': {'group': 'trend', 'name': 'DEMA (20)'},
    'MACD': {'group': 'trend', 'name': 'MACD + Crossover'},
    'SAR': {'group': 'trend', 'name': 'Parabolic SAR'},
    'RSI': {'group': 'momentum', 'name': 'RSI (14)'},
    'STOCH': {'group': 'momentum', 'name': 'Stochastic %K/%D'},
    'STOCHRSI': {'group': 'momentum', 'name': 'Stochastic RSI'},
    'ROC': {'group': 'momentum', 'name': 'ROC (10)'},
    'MOM': {'group': 'momentum', 'name': 'Momentum (10)'},
    'KAMA': {'group': 'momentum', 'name': 'KAMA (10)'},
    'CCI': {'group': 'oscillator', 'name': 'CCI (20)'},
    'WILLR': {'group': 'oscillator', 'name': 'Williams %R'},
    'ADX': {'group': 'oscillator', 'name': 'ADX (14)'},
    'ATR': {'group': 'oscillator', 'name': 'ATR (14)'},
    'BB': {'group': 'oscillator', 'name': 'Bollinger Bands'},
    'ULTOSC': {'group': 'oscillator', 'name': 'Ultimate Oscillator'},
    'OBV': {'group': 'volume', 'name': 'OBV'},
    'MFI': {'group': 'volume', 'name': 'MFI (14)'},
    'CMF': {'group': 'volume', 'name': 'CMF (20)'},
    'AD': {'group': 'volume', 'name': 'A/D Line'},
    'VWAP': {'group': 'volume', 'name': 'VWAP'},
    'PPO': {'group': 'volume', 'name': 'PPO'},
    'TRIX': {'group': 'volume', 'name': 'TRIX'},
}

# ============================================================
# HÀM TÍNH CHỈ BÁO (26 CHỈ BÁO)
# ============================================================

def calculate_indicators(df, selected_indicators):
    """Tính toán tất cả chỉ báo được chọn"""
    results = {}
    c = df['close'].astype(float)
    h = df['high'].astype(float)
    l = df['low'].astype(float)
    o = df['open'].astype(float)
    v = df['volume'].astype(float)
    n = len(df)
    
    # === TREND ===
    if 'SMA' in selected_indicators:
        for p in [5, 10, 20, 50, 100, 200]:
            if n >= p:
                results[f'SMA_{p}'] = c.rolling(p).mean().round(2)
    
    if 'EMA' in selected_indicators:
        for p in [12, 26, 50]:
            if n >= p:
                results[f'EMA_{p}'] = c.ewm(span=p, adjust=False).mean().round(2)
    
    if 'WMA' in selected_indicators:
        for p in [10, 20]:
            if n >= p:
                weights = np.arange(1, p + 1)
                results[f'WMA_{p}'] = c.rolling(p).apply(
                    lambda x: np.dot(x, weights) / weights.sum(), raw=True
                ).round(2)
    
    if 'TEMA' in selected_indicators and n >= 20:
        e1 = c.ewm(span=20, adjust=False).mean()
        e2 = e1.ewm(span=20, adjust=False).mean()
        e3 = e2.ewm(span=20, adjust=False).mean()
        results['TEMA_20'] = (3 * e1 - 3 * e2 + e3).round(2)
    
    if 'DEMA' in selected_indicators and n >= 20:
        e1 = c.ewm(span=20, adjust=False).mean()
        e2 = e1.ewm(span=20, adjust=False).mean()
        results['DEMA_20'] = (2 * e1 - e2).round(2)
    
    if 'MACD' in selected_indicators and n >= 26:
        e12 = c.ewm(span=12, adjust=False).mean()
        e26 = c.ewm(span=26, adjust=False).mean()
        macd = e12 - e26
        sig = macd.ewm(span=9, adjust=False).mean()
        hist = macd - sig
        
        results['MACD'] = macd.round(3)
        results['MACD_Signal'] = sig.round(3)
        results['MACD_Hist'] = hist.round(3)
        
        h_prev = hist.shift(1)
        cross = pd.Series('', index=c.index)
        cross[(h_prev < 0) & (hist > 0)] = 'CẮT_LÊN'
        cross[(h_prev > 0) & (hist < 0)] = 'CẮT_XUỐNG'
        results['MACD_Cross'] = cross
    
    if 'SAR' in selected_indicators and n >= 5:
        sar = pd.Series(index=c.index, dtype=float)
        sar.iloc[0] = l.iloc[0]
        af, af_max = 0.02, 0.2
        ep = h.iloc[0]
        trend = 1
        cur_af = af
        for i in range(1, n):
            if trend == 1:
                sar.iloc[i] = sar.iloc[i-1] + cur_af * (ep - sar.iloc[i-1])
                if h.iloc[i] > ep:
                    ep = h.iloc[i]
                    cur_af = min(cur_af + af, af_max)
                if l.iloc[i] < sar.iloc[i]:
                    trend = -1
                    sar.iloc[i] = ep
                    ep = l.iloc[i]
                    cur_af = af
            else:
                sar.iloc[i] = sar.iloc[i-1] + cur_af * (ep - sar.iloc[i-1])
                if l.iloc[i] < ep:
                    ep = l.iloc[i]
                    cur_af = min(cur_af + af, af_max)
                if h.iloc[i] > sar.iloc[i]:
                    trend = 1
                    sar.iloc[i] = ep
                    ep = h.iloc[i]
                    cur_af = af
        results['SAR'] = sar.round(2)
    
    # === MOMENTUM ===
    if 'RSI' in selected_indicators and n >= 14:
        delta = c.diff()
        gain = delta.where(delta > 0, 0).rolling(14).mean()
        loss = (-delta.where(delta < 0, 0)).rolling(14).mean()
        rs = gain / loss.replace(0, np.nan)
        results['RSI'] = (100 - 100 / (1 + rs)).round(2)
    
    if 'STOCH' in selected_indicators and n >= 14:
        low_14 = l.rolling(14).min()
        high_14 = h.rolling(14).max()
        denom = (high_14 - low_14).replace(0, np.nan)
        results['Stoch_K'] = (100 * (c - low_14) / denom).round(2)
        results['Stoch_D'] = results['Stoch_K'].rolling(3).mean().round(2)
    
    if 'STOCHRSI' in selected_indicators and n >= 28:
        delta = c.diff()
        gain = delta.where(delta > 0, 0).rolling(14).mean()
        loss = (-delta.where(delta < 0, 0)).rolling(14).mean()
        rsi = 100 - 100 / (1 + gain / loss.replace(0, np.nan))
        low_rsi = rsi.rolling(14).min()
        high_rsi = rsi.rolling(14).max()
        denom = (high_rsi - low_rsi).replace(0, np.nan)
        results['StochRSI'] = ((rsi - low_rsi) / denom * 100).round(2)
    
    if 'ROC' in selected_indicators and n >= 10:
        results['ROC'] = ((c - c.shift(10)) / c.shift(10) * 100).round(2)
    
    if 'MOM' in selected_indicators and n >= 10:
        results['Momentum'] = (c - c.shift(10)).round(2)
    
    if 'KAMA' in selected_indicators and n >= 10:
        change = abs(c - c.shift(10))
        volatility = abs(c.diff()).rolling(10).sum()
        er = change / volatility.replace(0, np.nan)
        fast = 2 / 3
        slow = 2 / 31
        sc = (er * (fast - slow) + slow) ** 2
        kama = pd.Series(index=c.index, dtype=float)
        kama.iloc[9] = c.iloc[9]
        for i in range(10, n):
            if pd.notna(sc.iloc[i]) and pd.notna(kama.iloc[i-1]):
                kama.iloc[i] = kama.iloc[i-1] + sc.iloc[i] * (c.iloc[i] - kama.iloc[i-1])
        results['KAMA'] = kama.round(2)
    
    # === OSCILLATOR ===
    if 'CCI' in selected_indicators and n >= 20:
        tp = (h + l + c) / 3
        sma_tp = tp.rolling(20).mean()
        mad = tp.rolling(20).apply(lambda x: np.abs(x - x.mean()).mean(), raw=True)
        results['CCI'] = ((tp - sma_tp) / (0.015 * mad)).round(2)
    
    if 'WILLR' in selected_indicators and n >= 14:
        high_14 = h.rolling(14).max()
        low_14 = l.rolling(14).min()
        denom = (high_14 - low_14).replace(0, np.nan)
        results['Williams_R'] = (-100 * (high_14 - c) / denom).round(2)
    
    if 'ADX' in selected_indicators and n >= 28:
        prev_c = c.shift(1)
        tr = pd.concat([h - l, abs(h - prev_c), abs(l - prev_c)], axis=1).max(axis=1)
        pdm = h.diff().clip(lower=0)
        mdm = (-l.diff()).clip(lower=0)
        pdm = pdm.where(pdm > mdm, 0)
        mdm = mdm.where(mdm > pdm, 0)
        atr = tr.ewm(span=14, adjust=False).mean()
        pdi = 100 * pdm.ewm(span=14, adjust=False).mean() / atr
        mdi = 100 * mdm.ewm(span=14, adjust=False).mean() / atr
        dx = 100 * abs(pdi - mdi) / (pdi + mdi).replace(0, np.nan)
        results['ADX'] = dx.ewm(span=14, adjust=False).mean().round(2)
        results['Plus_DI'] = pdi.round(2)
        results['Minus_DI'] = mdi.round(2)
    
    if 'ATR' in selected_indicators and n >= 14:
        prev_c = c.shift(1)
        tr = pd.concat([h - l, abs(h - prev_c), abs(l - prev_c)], axis=1).max(axis=1)
        results['ATR'] = tr.ewm(span=14, adjust=False).mean().round(2)
    
    if 'BB' in selected_indicators and n >= 20:
        sma = c.rolling(20).mean()
        std = c.rolling(20).std()
        results['BB_Upper'] = (sma + 2 * std).round(2)
        results['BB_Middle'] = sma.round(2)
        results['BB_Lower'] = (sma - 2 * std).round(2)
    
    if 'ULTOSC' in selected_indicators and n >= 28:
        prev_c = c.shift(1)
        bp = c - pd.concat([l, prev_c], axis=1).min(axis=1)
        tr = pd.concat([h - l, abs(h - prev_c), abs(l - prev_c)], axis=1).max(axis=1)
        avg7 = bp.rolling(7).sum() / tr.rolling(7).sum()
        avg14 = bp.rolling(14).sum() / tr.rolling(14).sum()
        avg28 = bp.rolling(28).sum() / tr.rolling(28).sum()
        results['UltOsc'] = (100 * (4 * avg7 + 2 * avg14 + avg28) / 7).round(2)
    
    # === VOLUME ===
    if 'OBV' in selected_indicators:
        obv = pd.Series(index=c.index, dtype=float)
        obv.iloc[0] = v.iloc[0]
        for i in range(1, n):
            if c.iloc[i] > c.iloc[i-1]:
                obv.iloc[i] = obv.iloc[i-1] + v.iloc[i]
            elif c.iloc[i] < c.iloc[i-1]:
                obv.iloc[i] = obv.iloc[i-1] - v.iloc[i]
            else:
                obv.iloc[i] = obv.iloc[i-1]
        results['OBV'] = obv.round(0)
    
    if 'MFI' in selected_indicators and n >= 14:
        tp = (h + l + c) / 3
        mf = tp * v
        pmf = mf.where(tp > tp.shift(1), 0).rolling(14).sum()
        nmf = mf.where(tp < tp.shift(1), 0).rolling(14).sum()
        mfr = pmf / nmf.replace(0, np.nan)
        results['MFI'] = (100 - 100 / (1 + mfr)).round(2)
    
    if 'CMF' in selected_indicators and n >= 20:
        denom = (h - l).replace(0, np.nan)
        mfv = ((c - l) - (h - c)) / denom * v
        results['CMF'] = (mfv.rolling(20).sum() / v.rolling(20).sum()).round(4)
    
    if 'AD' in selected_indicators:
        denom = (h - l).replace(0, np.nan)
        clv = ((c - l) - (h - c)) / denom
        results['AD'] = (clv.fillna(0) * v).cumsum().round(0)
    
    if 'VWAP' in selected_indicators:
        tp = (h + l + c) / 3
        results['VWAP'] = ((tp * v).cumsum() / v.cumsum()).round(2)
    
    if 'PPO' in selected_indicators and n >= 26:
        e12 = c.ewm(span=12, adjust=False).mean()
        e26 = c.ewm(span=26, adjust=False).mean()
        results['PPO'] = ((e12 - e26) / e26 * 100).round(2)
    
    if 'TRIX' in selected_indicators and n >= 45:
        e1 = c.ewm(span=15, adjust=False).mean()
        e2 = e1.ewm(span=15, adjust=False).mean()
        e3 = e2.ewm(span=15, adjust=False).mean()
        results['TRIX'] = (e3.pct_change() * 100).round(4)
    
    return results
PHẦN 2/2: app.py (Tiếp theo → cuối file)
Copy# ============================================================
# HÀM PHÂN TÍCH HÀNH ĐỘNG (GIỐNG COLAB)
# ============================================================

def get_action_conclusion(row, market_type='neutral'):
    """
    Đưa ra kết luận hành động - GIỐNG Y HỆT BẢN COLAB
    Trả về: action, confidence, reasons, conditions, ref_prices
    """
    close = row.get('close', 0)
    rsi = row.get('RSI')
    stoch_k = row.get('Stoch_K')
    stoch_d = row.get('Stoch_D')
    macd_hist = row.get('MACD_Hist')
    macd_cross = row.get('MACD_Cross', '')
    mfi = row.get('MFI')
    bb_upper = row.get('BB_Upper')
    bb_lower = row.get('BB_Lower')
    bb_middle = row.get('BB_Middle')
    sma_20 = row.get('SMA_20')
    sma_50 = row.get('SMA_50')
    vol_ratio = row.get('Vol_Ratio')
    
    buy_points = 0
    sell_points = 0
    reasons_buy = []
    reasons_sell = []
    reasons_wait = []
    conditions_buy = []
    conditions_sell = []
    
    # === MACD CROSSOVER ===
    if macd_cross == 'CẮT_LÊN':
        buy_points += 5
        reasons_buy.append("✅ MACD vừa cắt lên (tín hiệu đảo chiều tăng)")
    elif macd_cross == 'CẮT_XUỐNG':
        sell_points += 5
        reasons_sell.append("✅ MACD vừa cắt xuống (tín hiệu đảo chiều giảm)")
    elif pd.notna(macd_hist):
        if macd_hist > 0:
            buy_points += 1
            reasons_buy.append(f"MACD Histogram dương ({macd_hist:.3f})")
        else:
            sell_points += 1
            reasons_sell.append(f"MACD Histogram âm ({macd_hist:.3f})")
            conditions_buy.append(f"Chờ MACD Histogram chuyển dương (hiện: {macd_hist:.3f})")
    
    # === RSI ===
    if pd.notna(rsi):
        if rsi < 30:
            buy_points += 3
            reasons_buy.append(f"✅ RSI = {rsi:.1f} < 30 (quá bán - vùng mua)")
        elif rsi < 40:
            buy_points += 1
            reasons_buy.append(f"RSI = {rsi:.1f} đang thấp")
            conditions_buy.append(f"Chờ RSI < 30 hoặc bắt đầu tăng")
        elif rsi > 70:
            sell_points += 3
            reasons_sell.append(f"✅ RSI = {rsi:.1f} > 70 (quá mua - vùng bán)")
        elif rsi > 60:
            sell_points += 1
            reasons_sell.append(f"RSI = {rsi:.1f} đang cao")
        else:
            reasons_wait.append(f"RSI = {rsi:.1f} ở vùng trung tính (30-70)")
    
    # === STOCHASTIC ===
    if pd.notna(stoch_k):
        if stoch_k < 20:
            buy_points += 2
            reasons_buy.append(f"✅ Stoch %K = {stoch_k:.1f} < 20 (quá bán)")
        elif stoch_k < 30:
            buy_points += 1
            conditions_buy.append(f"Chờ Stoch < 20 (hiện: {stoch_k:.1f})")
        elif stoch_k > 80:
            sell_points += 2
            reasons_sell.append(f"✅ Stoch %K = {stoch_k:.1f} > 80 (quá mua)")
        elif stoch_k > 70:
            sell_points += 1
    
    # === MFI ===
    if pd.notna(mfi):
        if mfi < 20:
            buy_points += 2
            reasons_buy.append(f"✅ MFI = {mfi:.1f} < 20 (tiền đang chảy vào)")
        elif mfi > 80:
            sell_points += 2
            reasons_sell.append(f"✅ MFI = {mfi:.1f} > 80 (tiền đang chảy ra)")
    
    # === BOLLINGER BANDS ===
    if pd.notna(bb_lower) and pd.notna(bb_upper) and close > 0:
        bb_pos = (close - bb_lower) / (bb_upper - bb_lower) * 100 if (bb_upper - bb_lower) > 0 else 50
        if close < bb_lower:
            buy_points += 2
            reasons_buy.append(f"✅ Giá ({close:,.0f}) dưới BB Lower ({bb_lower:,.0f})")
        elif bb_pos < 20:
            buy_points += 1
            reasons_buy.append(f"Giá gần BB Lower (vị trí: {bb_pos:.0f}%)")
        elif close > bb_upper:
            sell_points += 2
            reasons_sell.append(f"✅ Giá ({close:,.0f}) trên BB Upper ({bb_upper:,.0f})")
        elif bb_pos > 80:
            sell_points += 1
            reasons_sell.append(f"Giá gần BB Upper (vị trí: {bb_pos:.0f}%)")
    
    # === TREND (SMA) ===
    if pd.notna(sma_20) and pd.notna(sma_50) and close > 0:
        if close > sma_20 > sma_50:
            buy_points += 2
            reasons_buy.append("✅ Uptrend: Giá > SMA20 > SMA50")
        elif close < sma_20 < sma_50:
            sell_points += 2
            reasons_sell.append("✅ Downtrend: Giá < SMA20 < SMA50")
        else:
            reasons_wait.append("Xu hướng không rõ ràng")
    
    # === VOLUME ===
    if pd.notna(vol_ratio) and vol_ratio > 2:
        reasons_wait.append(f"⚡ Khối lượng đột biến x{vol_ratio:.1f} - cần chú ý")
    
    # === TỔNG HỢP ===
    net_score = buy_points - sell_points
    
    market_bonus = 0
    market_note = ""
    if market_type == 'bullish':
        market_bonus = 2
        market_note = "Thị trường đang tăng mạnh → Hỗ trợ mua"
    elif market_type == 'slightly_bullish':
        market_bonus = 1
        market_note = "Thị trường đang tăng → Thuận lợi cho mua"
    elif market_type == 'bearish':
        market_bonus = -2
        market_note = "Thị trường đang giảm mạnh → Hạn chế mua"
    elif market_type == 'slightly_bearish':
        market_bonus = -1
        market_note = "Thị trường đang giảm → Thận trọng khi mua"
    
    adjusted_score = net_score + market_bonus
    
    # === KẾT LUẬN ===
    if macd_cross == 'CẮT_LÊN' and buy_points >= 5:
        action = '🟢 MUA NGAY'
        confidence = 'CAO'
        main_reason = "MACD vừa cắt lên + nhiều chỉ báo ủng hộ"
    elif macd_cross == 'CẮT_XUỐNG' and sell_points >= 5:
        action = '🔴 BÁN NGAY'
        confidence = 'CAO'
        main_reason = "MACD vừa cắt xuống + nhiều chỉ báo ủng hộ"
    elif adjusted_score >= 6:
        action = '🟢 MUA'
        confidence = 'CAO'
        main_reason = "Nhiều chỉ báo cho tín hiệu mua"
    elif adjusted_score >= 3:
        action = '🟢 CÂN NHẮC MUA'
        confidence = 'TRUNG BÌNH'
        main_reason = "Một số chỉ báo cho tín hiệu mua"
    elif adjusted_score <= -6:
        action = '🔴 BÁN'
        confidence = 'CAO'
        main_reason = "Nhiều chỉ báo cho tín hiệu bán"
    elif adjusted_score <= -3:
        action = '🔴 CÂN NHẮC BÁN'
        confidence = 'TRUNG BÌNH'
        main_reason = "Một số chỉ báo cho tín hiệu bán"
    elif buy_points > 0 and sell_points == 0:
        action = '🟡 CHỜ MUA'
        confidence = 'THẤP'
        main_reason = "Có tín hiệu tích cực nhưng chưa đủ mạnh"
    elif sell_points > 0 and buy_points == 0:
        action = '🟡 CHỜ BÁN'
        confidence = 'THẤP'
        main_reason = "Có tín hiệu tiêu cực nhưng chưa đủ mạnh"
    else:
        action = '⚪ KHÔNG HÀNH ĐỘNG'
        confidence = 'N/A'
        main_reason = "Không có tín hiệu rõ ràng"
    
    # Điều kiện theo dõi
    if 'CHỜ' in action or 'CÂN NHẮC' in action:
        if not conditions_buy and buy_points > sell_points:
            conditions_buy.append("Chờ MACD cắt lên để xác nhận")
        if not conditions_sell and sell_points > buy_points:
            conditions_sell.append("Chờ MACD cắt xuống để xác nhận")
    
    # Giá tham khảo
    ref_prices = {}
    if pd.notna(bb_lower):
        ref_prices['Hỗ trợ (BB Lower)'] = bb_lower
    if pd.notna(bb_upper):
        ref_prices['Kháng cự (BB Upper)'] = bb_upper
    if pd.notna(sma_20):
        ref_prices['SMA 20'] = sma_20
    if pd.notna(sma_50):
        ref_prices['SMA 50'] = sma_50
    
    return {
        'action': action,
        'confidence': confidence,
        'main_reason': main_reason,
        'buy_points': buy_points,
        'sell_points': sell_points,
        'net_score': net_score,
        'adjusted_score': adjusted_score,
        'reasons_buy': reasons_buy,
        'reasons_sell': reasons_sell,
        'reasons_wait': reasons_wait,
        'conditions_buy': conditions_buy,
        'conditions_sell': conditions_sell,
        'market_note': market_note,
        'ref_prices': ref_prices,
    }

# ============================================================
# DỰ BÁO ĐA KHUNG THỜI GIAN (T1-T5, W1-W4, M1-M3)
# ============================================================

def detect_market_phase(df, lookback=10):
    """Xác định pha thị trường"""
    if len(df) < lookback:
        return 'KHÔNG_ĐỦ_DỮ_LIỆU', 0
    
    recent = df.tail(lookback).copy()
    close = recent['close'].values
    high = recent['high'].values
    low = recent['low'].values
    
    current_close = close[-1]
    min_low = np.min(low)
    max_high = np.max(high)
    
    price_range = max_high - min_low
    position_pct = (current_close - min_low) / price_range * 100 if price_range > 0 else 50
    
    trend_3d = (close[-1] - close[-3]) / close[-3] * 100 if len(close) >= 3 else 0
    
    rsi = df['RSI'].iloc[-1] if 'RSI' in df.columns and pd.notna(df['RSI'].iloc[-1]) else 50
    stoch_k = df['Stoch_K'].iloc[-1] if 'Stoch_K' in df.columns and pd.notna(df['Stoch_K'].iloc[-1]) else 50
    
    macd_hist_trend = 0
    if 'MACD_Hist' in df.columns and len(df) >= 3:
        hists = df['MACD_Hist'].tail(3).values
        if all(pd.notna(hists)):
            macd_hist_trend = 1 if hists[-1] > hists[-2] else -1
    
    confidence = 50
    
    if position_pct < 25 and rsi < 35 and stoch_k < 30:
        if trend_3d > -0.5 or macd_hist_trend > 0:
            phase = 'ĐÁY'
            confidence = min(90, 50 + (35 - rsi) + (30 - stoch_k) / 2)
        else:
            phase = 'GIẢM'
            confidence = 60
    elif position_pct > 75 and rsi > 65 and stoch_k > 70:
        if trend_3d < 0.5 or macd_hist_trend < 0:
            phase = 'ĐỈNH'
            confidence = min(90, 50 + (rsi - 65) + (stoch_k - 70) / 2)
        else:
            phase = 'TĂNG'
            confidence = 60
    elif trend_3d > 1 and 45 < rsi < 70:
        phase = 'TĂNG'
        confidence = min(80, 50 + trend_3d * 5)
    elif trend_3d < -1 and 30 < rsi < 55:
        phase = 'GIẢM'
        confidence = min(80, 50 + abs(trend_3d) * 5)
    else:
        phase = 'TÍCH_LŨY'
        confidence = 50
    
    return phase, round(confidence, 1)


def forecast_multi_timeframe(df, symbol):
    """Dự báo T1-T5, W1-W4, M1-M3"""
    if len(df) < 20:
        return None
    
    df = df.sort_values('time').reset_index(drop=True)
    current = df.iloc[-1]
    close = float(current['close'])
    
    # Chỉ báo
    rsi = current.get('RSI', 50)
    rsi = rsi if pd.notna(rsi) else 50
    macd_hist = current.get('MACD_Hist', 0)
    macd_hist = macd_hist if pd.notna(macd_hist) else 0
    stoch_k = current.get('Stoch_K', 50)
    stoch_k = stoch_k if pd.notna(stoch_k) else 50
    
    bb_upper = current.get('BB_Upper', close * 1.05)
    bb_lower = current.get('BB_Lower', close * 0.95)
    sma_20 = current.get('SMA_20', close)
    
    atr = current.get('ATR', close * 0.02)
    atr = atr if pd.notna(atr) else close * 0.02
    atr_pct = (atr / close) * 100
    
    phase, phase_conf = detect_market_phase(df)
    
    # Momentum
    closes = df['close'].tail(20).values
    mom_1d = (closes[-1] - closes[-2]) / closes[-2] * 100 if len(closes) >= 2 else 0
    mom_3d = (closes[-1] - closes[-4]) / closes[-4] * 100 if len(closes) >= 4 else 0
    mom_5d = (closes[-1] - closes[-6]) / closes[-6] * 100 if len(closes) >= 6 else 0
    mom_10d = (closes[-1] - closes[-11]) / closes[-11] * 100 if len(closes) >= 11 else 0
    
    avg_momentum = mom_1d * 0.3 + mom_3d * 0.3 + mom_5d * 0.2 + mom_10d * 0.2
    
    # MACD trend
    macd_trend = 0
    if 'MACD_Hist' in df.columns and len(df) >= 3:
        hists = df['MACD_Hist'].tail(3).values
        if all(pd.notna(hists)):
            macd_trend = 1 if hists[-1] > hists[-2] else -1
    
    max_daily = min(atr_pct * 1.5, 3.0)
    
    # Tính daily change theo pha
    if phase == 'ĐÁY':
        base = 0.5 + (30 - min(rsi, 30)) * 0.05
        if macd_trend > 0:
            base += 0.3
        daily_change = min(base, max_daily)
        decay = 0.85
    elif phase == 'ĐỈNH':
        base = -0.5 - (min(rsi, 85) - 70) * 0.05
        if macd_trend < 0:
            base -= 0.3
        daily_change = max(base, -max_daily)
        decay = 0.85
    elif phase == 'TĂNG':
        if rsi > 65:
            daily_change = avg_momentum * 0.2
            decay = 0.7
        else:
            daily_change = avg_momentum * 0.4
            decay = 0.9
    elif phase == 'GIẢM':
        if rsi < 35:
            daily_change = avg_momentum * 0.2
            decay = 0.7
        else:
            daily_change = avg_momentum * 0.4
            decay = 0.9
    else:
        # TÍCH_LŨY
        bullish = sum([
            1 if rsi < 45 else 0,
            2 if macd_trend > 0 else 0,
            1 if macd_hist > 0 else 0,
            1 if stoch_k < 40 else 0,
        ])
        bearish = sum([
            1 if rsi > 55 else 0,
            2 if macd_trend < 0 else 0,
            1 if macd_hist < 0 else 0,
            1 if stoch_k > 60 else 0,
        ])
        
        net = bullish - bearish
        if net >= 2:
            daily_change = max(0.2, avg_momentum * 0.3)
        elif net <= -2:
            daily_change = min(-0.2, avg_momentum * 0.3)
        else:
            daily_change = avg_momentum * 0.2
        decay = 0.8
    
    daily_change = max(min(daily_change, max_daily), -max_daily)
    
    forecasts = {'symbol': symbol, 'gia_T0': close, 'pha': phase}
    
    # Ngày T1-T5
    cumulative = 0
    for d in range(1, 6):
        day_change = daily_change * (decay ** (d - 1))
        cumulative += day_change
        forecasts[f'gia_T{d}'] = round(close * (1 + cumulative / 100), 2)
        forecasts[f'pct_T{d}'] = round(cumulative, 2)
    
    # Tuần W1-W4
    weekly_change = daily_change * 5 * 0.8
    cumulative = 0
    for w in range(1, 5):
        week_change = weekly_change * (0.7 ** (w - 1))
        cumulative += week_change
        forecasts[f'gia_W{w}'] = round(close * (1 + cumulative / 100), 2)
        forecasts[f'pct_W{w}'] = round(cumulative, 2)
    
    # Tháng M1-M3
    monthly_change = daily_change * 22 * 0.6
    cumulative = 0
    for m in range(1, 4):
        month_change = monthly_change * (0.6 ** (m - 1))
        cumulative += month_change
        forecasts[f'gia_M{m}'] = round(close * (1 + cumulative / 100), 2)
        forecasts[f'pct_M{m}'] = round(cumulative, 2)
    
    # Hành động
    if phase == 'ĐÁY':
        forecasts['hanh_dong'] = 'MUA'
        forecasts['rui_ro_mua_duoi'] = 'THẤP'
    elif phase == 'ĐỈNH':
        forecasts['hanh_dong'] = 'BÁN'
        forecasts['rui_ro_mua_duoi'] = 'RẤT_CAO'
    elif phase == 'TĂNG' and rsi < 65:
        forecasts['hanh_dong'] = 'GIỮ/MUA_THÊM'
        forecasts['rui_ro_mua_duoi'] = 'TRUNG_BÌNH'
    elif phase == 'GIẢM' and rsi > 35:
        forecasts['hanh_dong'] = 'CHỜ'
        forecasts['rui_ro_mua_duoi'] = 'TRUNG_BÌNH'
    else:
        forecasts['hanh_dong'] = 'THEO_DÕI'
        forecasts['rui_ro_mua_duoi'] = 'TRUNG_BÌNH'
    
    # Giá mua/bán
    bb_l = float(bb_lower) if pd.notna(bb_lower) else close * 0.95
    bb_u = float(bb_upper) if pd.notna(bb_upper) else close * 1.05
    support = min(bb_l, df['low'].tail(10).min())
    resist = max(bb_u, df['high'].tail(10).max())
    
    forecasts['gia_mua'] = round(support, 2)
    forecasts['gia_ban'] = round(resist, 2)
    forecasts['cat_lo'] = round(support * 0.97, 2)
    
    forecasts['rsi'] = round(rsi, 1)
    forecasts['stoch_k'] = round(stoch_k, 1)
    forecasts['momentum'] = round(avg_momentum, 2)
    
    return forecasts

# ============================================================
# LẤY DỮ LIỆU
# ============================================================

def get_stock_data(symbol, start_date, end_date):
    """Lấy dữ liệu từ vnstock"""
    try:
        from vnstock import Vnstock
        stock = Vnstock().stock(symbol=symbol, source='VCI')
        df = stock.quote.history(start=start_date, end=end_date, interval='1D')
        
        if df is not None and len(df) > 0:
            df['time'] = pd.to_datetime(df['time'])
            for col in ['open', 'high', 'low', 'close', 'volume']:
                df[col] = df[col].astype(float)
            return df
    except Exception as e:
        print(f"Lỗi {symbol}: {e}")
    return None


def get_market_indices():
    """Lấy chỉ số thị trường"""
    indices = {'VNINDEX': 'VN-Index', 'VN30': 'VN30', 'HNX30': 'HNX30'}
    data = []
    
    for code, name in indices.items():
        try:
            from vnstock import Vnstock
            stock = Vnstock().stock(symbol=code, source='VCI')
            df = stock.quote.history(
                start=(datetime.now() - timedelta(days=400)).strftime('%Y-%m-%d'),
                end=datetime.now().strftime('%Y-%m-%d'),
                interval='1D'
            )
            
            if df is not None and len(df) > 0:
                df['time'] = pd.to_datetime(df['time'])
                df = df.sort_values('time', ascending=False)
                cur = df.iloc[0]['close']
                
                def pct(d):
                    if len(df) > d and df.iloc[d]['close'] > 0:
                        return round((cur - df.iloc[d]['close']) / df.iloc[d]['close'] * 100, 2)
                    return None
                
                data.append({
                    'Chỉ số': name,
                    'Giá': round(cur, 2),
                    'Ngày': df.iloc[0]['time'].strftime('%d/%m/%Y'),
                    'D%': pct(1),
                    'W%': pct(5),
                    'M%': pct(22),
                    'Q%': pct(66),
                    'Y%': pct(252),
                })
        except:
            pass
    
    return pd.DataFrame(data) if data else None


def analyze_market(df_m):
    """Phân tích thị trường"""
    if df_m is None or len(df_m) == 0:
        return "Không có dữ liệu", "neutral"
    
    vni = df_m[df_m['Chỉ số'] == 'VN-Index']
    if len(vni) == 0:
        return "Không có VN-Index", "neutral"
    
    d = vni.iloc[0].get('D%', 0) or 0
    w = vni.iloc[0].get('W%', 0) or 0
    
    if d > 1 and w > 3:
        return "🟢 TĂNG MẠNH", "bullish"
    elif d > 0 and w > 0:
        return "🟢 TĂNG", "slightly_bullish"
    elif d < -1 and w < -3:
        return "🔴 GIẢM MẠNH", "bearish"
    elif d < 0 and w < 0:
        return "🔴 GIẢM", "slightly_bearish"
    else:
        return "🟡 TÍCH LŨY", "neutral"

# ============================================================
# XUẤT FILE WORD (GIỐNG Y HỆT COLAB)
# ============================================================

def export_word_full(summary_df, forecast_df, action_details, market_df, market_ctx, 
                     start_date, end_date, output_path):
    """Xuất file Word - GIỐNG Y HỆT BẢN COLAB (4 phần)"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # === TIÊU ĐỀ ===
        title = doc.add_heading('BÁO CÁO PHÂN TÍCH KỸ THUẬT', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f'Ngày: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
        doc.add_paragraph(f'Thời gian: {start_date} → {end_date}')
        doc.add_paragraph(f'Số mã: {len(action_details)}/{len(action_details)}')
        
        # === I. CHỈ SỐ THỊ TRƯỜNG ===
        if market_df is not None and len(market_df) > 0:
            doc.add_heading('I. CHỈ SỐ THỊ TRƯỜNG', level=1)
            
            table = doc.add_table(rows=1, cols=len(market_df.columns))
            table.style = 'Table Grid'
            
            for i, col in enumerate(market_df.columns):
                table.rows[0].cells[i].text = str(col)
            
            for _, row in market_df.iterrows():
                cells = table.add_row().cells
                for i, col in enumerate(market_df.columns):
                    val = row[col]
                    if isinstance(val, float) and pd.notna(val):
                        cells[i].text = f'{val:.2f}'
                    else:
                        cells[i].text = str(val) if pd.notna(val) else 'N/A'
            
            doc.add_paragraph('')
            doc.add_paragraph(f'📊 NHẬN ĐỊNH: {market_ctx}')
        
        # === II. TÓM TẮT HÀNH ĐỘNG ===
        doc.add_heading('II. TÓM TẮT HÀNH ĐỘNG', level=1)
        
        actions_group = {}
        for item in action_details:
            action = item['conclusion']['action']
            if action not in actions_group:
                actions_group[action] = []
            actions_group[action].append(item['symbol'])
        
        for action, symbols in actions_group.items():
            doc.add_paragraph(f"{action}: {len(symbols)} mã - {', '.join(symbols)}")
        
        # === III. KHUYẾN NGHỊ CHI TIẾT TỪNG MÃ ===
        doc.add_heading('III. KHUYẾN NGHỊ CHI TIẾT TỪNG MÃ', level=1)
        
        for item in action_details:
            sym = item['symbol']
            price = item['price']
            c = item['conclusion']
            
            doc.add_heading(f'{sym}: {price:,.0f}', level=2)
            
            # Bảng tóm tắt
            t = doc.add_table(rows=4, cols=2)
            t.style = 'Table Grid'
            t.rows[0].cells[0].text = 'HÀNH ĐỘNG'
            t.rows[0].cells[1].text = c['action']
            t.rows[1].cells[0].text = 'Độ tin cậy'
            t.rows[1].cells[1].text = c['confidence']
            t.rows[2].cells[0].text = 'Điểm MUA / BÁN'
            t.rows[2].cells[1].text = f"{c['buy_points']} / {c['sell_points']}"
            t.rows[3].cells[0].text = 'Lý do chính'
            t.rows[3].cells[1].text = c['main_reason']
            
            # Lý do MUA
            if c['reasons_buy']:
                doc.add_paragraph('✅ Lý do MUA:')
                for r in c['reasons_buy']:
                    doc.add_paragraph(f'   • {r}')
            
            # Lý do BÁN
            if c['reasons_sell']:
                doc.add_paragraph('❌ Lý do BÁN:')
                for r in c['reasons_sell']:
                    doc.add_paragraph(f'   • {r}')
            
            # Điều kiện MUA (cần theo dõi)
            if c['conditions_buy']:
                doc.add_paragraph('📋 Điều kiện MUA (cần theo dõi):')
                for r in c['conditions_buy']:
                    doc.add_paragraph(f'   • {r}')
            
            # Điều kiện BÁN (cần theo dõi)
            if c['conditions_sell']:
                doc.add_paragraph('📋 Điều kiện BÁN (cần theo dõi):')
                for r in c['conditions_sell']:
                    doc.add_paragraph(f'   • {r}')
            
            # Giá tham khảo
            if c['ref_prices']:
                doc.add_paragraph('💰 Giá tham khảo:')
                for name, val in c['ref_prices'].items():
                    doc.add_paragraph(f'   • {name}: {val:,.0f}')
            
            doc.add_paragraph('')
        
        # === IV. HƯỚNG DẪN ĐỌC BÁO CÁO ===
        doc.add_heading('IV. HƯỚNG DẪN ĐỌC BÁO CÁO', level=1)
        
        notes = [
            '🟢 MUA NGAY: Tín hiệu rất mạnh, có thể mua ngay',
            '🟢 MUA: Tín hiệu tốt, nên mua',
            '🟢 CÂN NHẮC MUA: Tín hiệu khá, cân nhắc mua',
            '🟡 CHỜ MUA: Chưa đủ tín hiệu, chờ xác nhận',
            '🔴 BÁN NGAY: Tín hiệu rất mạnh, nên bán ngay',
            '🔴 BÁN: Tín hiệu tốt, nên bán',
            '🔴 CÂN NHẮC BÁN: Tín hiệu khá, cân nhắc bán',
            '🟡 CHỜ BÁN: Chưa đủ tín hiệu, chờ xác nhận',
            '⚪ KHÔNG HÀNH ĐỘNG: Không có tín hiệu rõ ràng',
            '',
            'Độ tin cậy:',
            '• CAO: Nhiều chỉ báo đồng thuận',
            '• TRUNG BÌNH: Một số chỉ báo đồng thuận',
            '• THẤP: Ít chỉ báo, cần xác nhận thêm',
        ]
        for n in notes:
            doc.add_paragraph(n)
        
        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Lỗi xuất Word: {e}")
        return False

# ============================================================
# XUẤT FILE EXCEL
# ============================================================

def export_excel_full(all_data, summary_df, forecast_df, action_details, market_df, output_path):
    """Xuất file Excel tổng hợp"""
    try:
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            # Tổng hợp
            if summary_df is not None:
                summary_df.to_excel(writer, sheet_name='Tổng hợp', index=False)
            
            # Dự báo T0-T3
            if forecast_df is not None:
                # Chọn cột hiển thị
                cols_export = ['symbol', 'gia_T0', 'gia_T1', 'gia_T2', 'gia_T3', 'gia_T4', 'gia_T5',
                              'pct_T1', 'pct_T2', 'pct_T3', 'pct_T4', 'pct_T5',
                              'gia_W1', 'gia_W2', 'gia_W3', 'gia_W4',
                              'pct_W1', 'pct_W2', 'pct_W3', 'pct_W4',
                              'gia_M1', 'gia_M2', 'gia_M3',
                              'pct_M1', 'pct_M2', 'pct_M3',
                              'pha', 'hanh_dong', 'rui_ro_mua_duoi',
                              'gia_mua', 'gia_ban', 'cat_lo']
                cols_exist = [c for c in cols_export if c in forecast_df.columns]
                forecast_df[cols_exist].to_excel(writer, sheet_name='Dự báo T0-T5-W-M', index=False)
            
            # Khuyến nghị chi tiết
            action_rows = []
            for item in action_details:
                c = item['conclusion']
                action_rows.append({
                    'Mã': item['symbol'],
                    'Giá': item['price'],
                    'HÀNH ĐỘNG': c['action'],
                    'Độ tin cậy': c['confidence'],
                    'Lý do chính': c['main_reason'],
                    'Điểm MUA': c['buy_points'],
                    'Điểm BÁN': c['sell_points'],
                    'Lý do MUA': ' | '.join(c['reasons_buy']),
                    'Lý do BÁN': ' | '.join(c['reasons_sell']),
                    'Điều kiện MUA': ' | '.join(c['conditions_buy']),
                    'Điều kiện BÁN': ' | '.join(c['conditions_sell']),
                    'Ghi chú TT': c['market_note'],
                })
            pd.DataFrame(action_rows).to_excel(writer, sheet_name='Khuyến nghị chi tiết', index=False)
            
            # Thị trường
            if market_df is not None:
                market_df.to_excel(writer, sheet_name='Thị trường', index=False)
            
            # Sheet từng mã
            for symbol, df in all_data.items():
                df_s = df.sort_values('time', ascending=False).copy()
                exp = pd.DataFrame()
                exp['Ngày'] = df_s['time'].dt.strftime('%d/%m/%Y')
                for c in ['open', 'high', 'low', 'close', 'volume']:
                    if c in df_s.columns:
                        exp[c] = df_s[c]
                for c in df_s.columns:
                    if c not in ['time', 'open', 'high', 'low', 'close', 'volume', 'Ma', 'Vol_SMA_20', 'GTGD']:
                        exp[c] = df_s[c]
                exp.to_excel(writer, sheet_name=symbol[:31], index=False)
        
        return True
    except Exception as e:
        print(f"Lỗi xuất Excel: {e}")
        return False

# ============================================================
# XUẤT FILE RIÊNG TỪNG MÃ
# ============================================================

def export_individual_files(all_data, output_folder, timestamp):
    """Xuất file Vietstock và MetaStock cho từng mã"""
    vs_folder = os.path.join(output_folder, 'Vietstock')
    ms_folder = os.path.join(output_folder, 'MetaStock')
    
    os.makedirs(vs_folder, exist_ok=True)
    os.makedirs(ms_folder, exist_ok=True)
    
    for symbol, df in all_data.items():
        try:
            df_s = df.sort_values('time', ascending=False)
            
            # Vietstock
            vs = pd.DataFrame()
            vs['Ngày'] = df_s['time'].dt.strftime('%d/%m/%Y')
            vs['Mã'] = symbol
            for c in ['open', 'high', 'low', 'close', 'volume']:
                vs[c] = df_s[c]
            for c in ['RSI', 'MACD', 'MACD_Hist', 'MACD_Cross', 'Stoch_K', 'MFI', 'SMA_20', 'BB_Upper', 'BB_Lower']:
                if c in df_s.columns:
                    vs[c] = df_s[c]
            vs.to_excel(f"{vs_folder}/{symbol}_{timestamp}.xlsx", index=False)
            
            # MetaStock
            ms = pd.DataFrame({
                'DATE': df_s['time'].dt.strftime('%Y%m%d'),
                'TICKER': symbol,
                'OPEN': df_s['open'],
                'HIGH': df_s['high'],
                'LOW': df_s['low'],
                'CLOSE': df_s['close'],
                'VOLUME': df_s['volume'].astype(int),
            })
            ms.to_csv(f"{ms_folder}/{symbol}_{timestamp}.csv", index=False)
        except:
            pass

# ============================================================
# TẠO BIỂU ĐỒ
# ============================================================

def create_chart(df, symbol):
    """Tạo biểu đồ giá và chỉ báo"""
    df = df.sort_values('time')
    
    fig = make_subplots(
        rows=4, cols=1,
        shared_xaxes=True,
        vertical_spacing=0.03,
        row_heights=[0.5, 0.15, 0.15, 0.2],
        subplot_titles=(f'{symbol} - Giá', 'RSI', 'MACD', 'Volume')
    )
    
    fig.add_trace(
        go.Candlestick(
            x=df['time'], open=df['open'], high=df['high'],
            low=df['low'], close=df['close'], name='Giá'
        ), row=1, col=1
    )
    
    if 'BB_Upper' in df.columns:
        fig.add_trace(go.Scatter(x=df['time'], y=df['BB_Upper'], name='BB Upper',
                                line=dict(color='gray', dash='dash')), row=1, col=1)
        fig.add_trace(go.Scatter(x=df['time'], y=df['BB_Lower'], name='BB Lower',
                                line=dict(color='gray', dash='dash')), row=1, col=1)
    
    if 'SMA_20' in df.columns:
        fig.add_trace(go.Scatter(x=df['time'], y=df['SMA_20'], name='SMA20',
                                line=dict(color='blue')), row=1, col=1)
    
    if 'RSI' in df.columns:
        fig.add_trace(go.Scatter(x=df['time'], y=df['RSI'], name='RSI',
                                line=dict(color='purple')), row=2, col=1)
        fig.add_hline(y=70, line_dash="dash", line_color="red", row=2, col=1)
        fig.add_hline(y=30, line_dash="dash", line_color="green", row=2, col=1)
    
    if 'MACD_Hist' in df.columns:
        colors = ['green' if v >= 0 else 'red' for v in df['MACD_Hist'].fillna(0)]
        fig.add_trace(go.Bar(x=df['time'], y=df['MACD_Hist'], name='MACD Hist',
                            marker_color=colors), row=3, col=1)
    
    fig.add_trace(go.Bar(x=df['time'], y=df['volume'], name='Volume',
                        marker_color='lightblue'), row=4, col=1)
    
    fig.update_layout(height=800, showlegend=False, xaxis_rangeslider_visible=False)
    
    return fig

# ============================================================
# HÀM CHẠY CHÍNH
# ============================================================

def run_analysis(symbols_text, days_back, trend_ind, momentum_ind, oscillator_ind, volume_ind, progress=gr.Progress()):
    """Hàm chạy phân tích chính"""
    
    # Gộp chỉ báo
    selected_indicators = trend_ind + momentum_ind + oscillator_ind + volume_ind
    
    # Parse symbols
    symbols = [s.strip().upper() for s in symbols_text.split(',') if s.strip()]
    if not symbols:
        return "Vui lòng nhập ít nhất 1 mã", None, None, None, None, None
    
    # Date range
    end_date = datetime.now().strftime('%Y-%m-%d')
    start_date = (datetime.now() - timedelta(days=int(days_back))).strftime('%Y-%m-%d')
    
    # Output folder
    date_folder = datetime.now().strftime('%Y%m%d')
    output_folder = f"./output/{date_folder}"
    os.makedirs(output_folder, exist_ok=True)
    
    # Results
    all_data = {}
    summary_rows = []
    forecast_rows = []
    action_details = []
    failed = []
    
    # Market
    progress(0.05, desc="Lấy chỉ số thị trường...")
    market_df = get_market_indices()
    market_ctx, market_type = analyze_market(market_df)
    
    # Process
    for i, symbol in enumerate(symbols):
        progress((i + 1) / len(symbols), desc=f"Đang xử lý {symbol}...")
        
        try:
            df = get_stock_data(symbol, start_date, end_date)
            
            if df is None or len(df) < 5:
                failed.append(symbol)
                continue
            
            # Indicators
            indicators = calculate_indicators(df, selected_indicators)
            for col, values in indicators.items():
                df[col] = values
            
            df['Vol_SMA_20'] = df['volume'].rolling(min(20, len(df))).mean()
            df['Vol_Ratio'] = df['volume'] / df['Vol_SMA_20']
            df['Thay_doi_pct'] = df['close'].pct_change() * 100
            
            all_data[symbol] = df.sort_values('time', ascending=False).reset_index(drop=True)
            
            latest = df.sort_values('time', ascending=False).iloc[0]
            
            # Action
            conclusion = get_action_conclusion(latest, market_type)
            
            summary_rows.append({
                'Mã': symbol,
                'Ngày': latest['time'].strftime('%d/%m/%Y'),
                'Giá': latest['close'],
                '%': round(latest.get('Thay_doi_pct', 0) or 0, 2),
                'KL': int(latest['volume']),
                'RSI': round(latest.get('RSI', 0) or 0, 1),
                'Stoch_K': round(latest.get('Stoch_K', 0) or 0, 1),
                'MACD_Cross': latest.get('MACD_Cross', ''),
                'Hành động': conclusion['action'],
                'Độ tin cậy': conclusion['confidence'],
            })
            
            # Forecast
            forecast = forecast_multi_timeframe(df, symbol)
            if forecast:
                forecast_rows.append(forecast)
            
            action_details.append({
                'symbol': symbol,
                'price': latest['close'],
                'conclusion': conclusion,
            })
            
        except Exception as e:
            failed.append(f"{symbol}: {str(e)[:30]}")
    
    if not all_data:
        return "Không lấy được dữ liệu!", None, None, None, None, None
    
    # DataFrames
    summary_df = pd.DataFrame(summary_rows)
    forecast_df = pd.DataFrame(forecast_rows) if forecast_rows else None
    
    # Export
    timestamp = datetime.now().strftime('%Y%m%d_%H%M')
    
    progress(0.9, desc="Xuất file...")
    
    excel_path = f"{output_folder}/BaoCao_{timestamp}.xlsx"
    export_excel_full(all_data, summary_df, forecast_df, action_details, market_df, excel_path)
    
    word_path = f"{output_folder}/BaoCao_{timestamp}.docx"
    export_word_full(summary_df, forecast_df, action_details, market_df, market_ctx,
                    start_date, end_date, word_path)
    
    export_individual_files(all_data, output_folder, timestamp)
    
    # ZIP
    zip_path = f"./output/BaoCao_{date_folder}_{timestamp}.zip"
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
        for root, dirs, files in os.walk(output_folder):
            for file in files:
                fp = os.path.join(root, file)
                zf.write(fp, os.path.relpath(fp, output_folder))
    
    # Results
    results_text = f"""
## ✅ HOÀN THÀNH PHÂN TÍCH

**Thông tin:**
- Số mã: {len(all_data)}/{len(symbols)}
- Thời gian: {start_date} → {end_date}
- Thị trường: {market_ctx}

**File đã tạo:**
- 📊 Excel: BaoCao_{timestamp}.xlsx
- 📝 Word: BaoCao_{timestamp}.docx (GIỐNG COLAB)
- 📁 Vietstock: {len(all_data)} file
- 📁 MetaStock: {len(all_data)} file
- 📦 ZIP: {os.path.basename(zip_path)}
"""
    
    if failed:
        results_text += f"\n**Lỗi:** {', '.join(failed)}"
    
    # Chart
    first_symbol = list(all_data.keys())[0]
    chart = create_chart(all_data[first_symbol], first_symbol)
    
    return results_text, summary_df, forecast_df, pd.DataFrame([{
        'Mã': a['symbol'],
        'Giá': a['price'],
        'Hành động': a['conclusion']['action'],
        'Độ tin cậy': a['conclusion']['confidence'],
    } for a in action_details]), chart, zip_path

# ============================================================
# GIAO DIỆN
# ============================================================

def create_ui():
    config = load_config()
    
    with gr.Blocks(title="Stock Analyzer VN", theme=gr.themes.Soft()) as app:
        
        gr.Markdown("""
        # 📈 HỆ THỐNG PHÂN TÍCH KỸ THUẬT CHỨNG KHOÁN VIỆT NAM
        **26 chỉ báo | Dự báo T1-T5, W1-W4, M1-M3 | Xuất Excel, Word (giống Colab)**
        """)
        
        with gr.Row():
            with gr.Column(scale=1):
                gr.Markdown("### ⚙️ Cấu hình")
                
                symbols_input = gr.Textbox(
                    label="Danh sách mã",
                    value=", ".join(config.get("watchlist", [])),
                    lines=2,
                )
                
                days_slider = gr.Slider(
                    label="Số ngày", minimum=30, maximum=365,
                    value=config.get("days_back", 90), step=1,
                )
                
                gr.Markdown("### 📊 Chỉ báo")
                
                with gr.Accordion("📈 Xu hướng", open=True):
                    trend_checks = gr.CheckboxGroup(
                        choices=['SMA', 'EMA', 'WMA', 'TEMA', 'DEMA', 'MACD', 'SAR'],
                        value=['SMA', 'EMA', 'MACD'],
                    )
                
                with gr.Accordion("⚡ Động lượng", open=True):
                    momentum_checks = gr.CheckboxGroup(
                        choices=['RSI', 'STOCH', 'STOCHRSI', 'ROC', 'MOM', 'KAMA'],
                        value=['RSI', 'STOCH'],
                    )
                
                with gr.Accordion("🔄 Dao động", open=True):
                    oscillator_checks = gr.CheckboxGroup(
                        choices=['CCI', 'WILLR', 'ADX', 'ATR', 'BB', 'ULTOSC'],
                        value=['BB', 'ATR'],
                    )
                
                with gr.Accordion("📊 Khối lượng", open=True):
                    volume_checks = gr.CheckboxGroup(
                        choices=['OBV', 'MFI', 'CMF', 'AD', 'VWAP', 'PPO', 'TRIX'],
                        value=['OBV', 'MFI'],
                    )
                
                run_btn = gr.Button("🚀 BẮT ĐẦU PHÂN TÍCH", variant="primary", size="lg")
            
            with gr.Column(scale=2):
                results_md = gr.Markdown("")
                
                with gr.Tabs():
                    with gr.Tab("📊 Tổng hợp"):
                        summary_table = gr.Dataframe()
                    with gr.Tab("🔮 Dự báo"):
                        forecast_table = gr.Dataframe()
                    with gr.Tab("🎯 Khuyến nghị"):
                        action_table = gr.Dataframe()
                    with gr.Tab("📈 Biểu đồ"):
                        chart_plot = gr.Plot()
                
                download_file = gr.File(label="📥 Tải ZIP")
        
        run_btn.click(
            fn=run_analysis,
            inputs=[symbols_input, days_slider, trend_checks, momentum_checks, oscillator_checks, volume_checks],
            outputs=[results_md, summary_table, forecast_table, action_table, chart_plot, download_file],
        )
        
        gr.Markdown("""
        ---
        ### 📖 Hướng dẫn
        | Ký hiệu | Ý nghĩa |
        |---------|---------|
        | 🟢 MUA NGAY/MUA | Tín hiệu mạnh |
        | 🟡 CHỜ MUA/BÁN | Chờ xác nhận |
        | 🔴 BÁN NGAY/BÁN | Nên bán |
        | ⚪ KHÔNG HÀNH ĐỘNG | Chưa rõ |
        """)
    
    return app

# ============================================================
# MAIN
# ============================================================

if __name__ == "__main__":
    print("="*60)
    print("🚀 KHỞI ĐỘNG HỆ THỐNG PHÂN TÍCH CHỨNG KHOÁN")
    print("="*60)
    
    os.makedirs("./output", exist_ok=True)
    config = load_config()
    print(f"✅ Đã tải cấu hình: {len(config.get('watchlist', []))} mã")
    
    app = create_ui()
    
    print("\n📌 Mở trình duyệt tại địa chỉ bên dưới:")
    app.launch(server_name="0.0.0.0", server_port=7860, share=False)
